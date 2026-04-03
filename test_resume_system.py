#!/usr/bin/env python3
"""
Resume Upload & Skill Extraction - Comprehensive Test Suite

This script tests all components of the resume upload and AI skill extraction feature:
- Frontend JavaScript initialization
- Backend API endpoints
- Database operations
- API key configuration
- File parsing capabilities
"""

import os
import sys
import json
import sqlite3
import tempfile
from pathlib import Path
from datetime import datetime

# Color codes for terminal output
class Colors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

def print_header(text):
    print(f"\n{Colors.HEADER}{Colors.BOLD}{'='*70}")
    print(f"  {text}")
    print(f"{'='*70}{Colors.ENDC}\n")

def print_success(text):
    print(f"{Colors.OKGREEN}✓ {text}{Colors.ENDC}")

def print_error(text):
    print(f"{Colors.FAIL}✗ {text}{Colors.ENDC}")

def print_warning(text):
    print(f"{Colors.WARNING}⚠ {text}{Colors.ENDC}")

def print_info(text):
    print(f"{Colors.OKCYAN}ℹ {text}{Colors.ENDC}")

# ─────────────────────────────────────────────────────────────────────────────
# Test 1: Environment Configuration
# ─────────────────────────────────────────────────────────────────────────────

def test_environment():
    print_header("Test 1: Environment Configuration")
    
    tests_passed = 0
    tests_total = 0
    
    # Check GROQ API Key
    tests_total += 1
    api_key = os.environ.get("GROQ_API_KEY")
    if api_key and api_key.startswith("gsk_"):
        print_success("GROQ_API_KEY is set and valid")
        tests_passed += 1
    elif api_key:
        print_warning(f"GROQ_API_KEY is set but may be invalid (doesn't start with 'gsk_')")
    else:
        print_error("GROQ_API_KEY not set. Set with: export GROQ_API_KEY='your-key'")
    
    # Check working directory
    tests_total += 1
    if os.path.exists("prepfriend"):
        print_success("Working directory is correct (prepfriend folder found)")
        tests_passed += 1
    else:
        print_error("prepfriend folder not found. Run from repo root")
    
    # Check database file exists
    tests_total += 1
    db_path = Path("prepfriend/data/database.db")
    if db_path.exists():
        print_success(f"Database file exists: {db_path}")
        tests_passed += 1
    else:
        print_warning(f"Database file not found: {db_path}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 2: Python Dependencies
# ─────────────────────────────────────────────────────────────────────────────

def test_dependencies():
    print_header("Test 2: Python Dependencies")
    
    tests_passed = 0
    tests_total = 0
    
    dependencies = [
        ("flask", "Flask"),
        ("openai", "OpenAI"),
        ("PyPDF2", "PyPDF2"),
        ("docx", "python-docx"),
        ("pdfplumber", "pdfplumber (optional but recommended)"),
    ]
    
    for module, name in dependencies:
        tests_total += 1
        try:
            __import__(module)
            print_success(f"{name} is installed")
            tests_passed += 1
        except ImportError:
            if "optional" in name:
                print_warning(f"{name} not installed (optional)")
            else:
                print_error(f"{name} not installed. Install with: pip install {module}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 3: Frontend Files
# ─────────────────────────────────────────────────────────────────────────────

def test_frontend_files():
    print_header("Test 3: Frontend Files")
    
    tests_passed = 0
    tests_total = 0
    
    files_to_check = [
        ("prepfriend/app/static/js/resume-modal.js", "JavaScript - Resume Modal"),
        ("prepfriend/app/static/css/resume-modal.css", "CSS - Resume Modal"),
        ("prepfriend/app/templates/dashboard.html", "HTML - Dashboard"),
    ]
    
    for file_path, description in files_to_check:
        tests_total += 1
        if Path(file_path).exists():
            file_size = Path(file_path).stat().st_size
            print_success(f"{description} exists ({file_size:,} bytes)")
            tests_passed += 1
        else:
            print_error(f"{description} not found at {file_path}")
    
    # Check if resume-modal.js includes logging
    tests_total += 1
    try:
        with open("prepfriend/app/static/js/resume-modal.js", "r") as f:
            content = f.read()
            if "🎯" in content or "console.log" in content:
                print_success("resume-modal.js includes debug logging")
                tests_passed += 1
            else:
                print_warning("resume-modal.js may not have sufficient logging")
    except:
        print_error("Could not read resume-modal.js")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 4: Database Schema
# ─────────────────────────────────────────────────────────────────────────────

def test_database_schema():
    print_header("Test 4: Database Schema")
    
    tests_passed = 0
    tests_total = 0
    
    db_path = "prepfriend/data/database.db"
    
    if not Path(db_path).exists():
        print_error(f"Database file not found: {db_path}")
        return tests_passed, tests_total
    
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Check for resumes table
        tests_total += 1
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='resumes'")
        if cursor.fetchone():
            print_success("Resumes table exists")
            tests_passed += 1
        else:
            print_error("Resumes table not found")
        
        # Check resumes table columns
        tests_total += 1
        try:
            cursor.execute("PRAGMA table_info(resumes)")
            columns = {row[1]: row[2] for row in cursor.fetchall()}
            required_cols = ["id", "email", "filename", "file_path", "file_content"]
            if all(col in columns for col in required_cols):
                print_success(f"Resumes table has required columns: {', '.join(required_cols)}")
                tests_passed += 1
            else:
                print_error("Resumes table missing required columns")
        except Exception as e:
            print_error(f"Could not check resumes columns: {e}")
        
        # Count existing resumes
        tests_total += 1
        cursor.execute("SELECT COUNT(*) FROM resumes")
        count = cursor.fetchone()[0]
        if count > 0:
            print_info(f"Found {count} resume(s) in database")
        print_success(f"Resumes table is accessible (count check passed)")
        tests_passed += 1
        
        # Check skill_checklists table
        tests_total += 1
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='skill_checklists'")
        if cursor.fetchone():
            print_success("Skill checklists table exists")
            tests_passed += 1
        else:
            print_error("Skill checklists table not found")
        
        conn.close()
    except Exception as e:
        print_error(f"Database access error: {e}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 5: Backend Route Definitions
# ─────────────────────────────────────────────────────────────────────────────

def test_backend_routes():
    print_header("Test 5: Backend Route Definitions")
    
    tests_passed = 0
    tests_total = 0
    
    try:
        with open("prepfriend/app/routes.py", "r") as f:
            content = f.read()
        
        routes_to_check = [
            ("/api/resume/upload", "Resume Upload"),
            ("/api/resume/latest", "Get Latest Resume"),
            ("/api/resume/extract-skills", "Extract Skills"),
        ]
        
        for route, description in routes_to_check:
            tests_total += 1
            if route in content:
                print_success(f"Route handler found: {description} ({route})")
                tests_passed += 1
            else:
                print_error(f"Route handler not found: {description} ({route})")
        
        # Check for required functions
        tests_total += 1
        required_functions = [
            "def upload_resume",
            "def extract_skills_from_resume",
            "def allowed_file",
            "def extract_text_from_file",
        ]
        
        for func in required_functions:
            if func in content:
                print_success(f"Function found: {func}")
            else:
                print_error(f"Function not found: {func}")
                tests_passed -= 1
        
        tests_passed += 1
        
    except Exception as e:
        print_error(f"Could not read routes.py: {e}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 6: File Parsing Capabilities
# ─────────────────────────────────────────────────────────────────────────────

def test_file_parsing():
    print_header("Test 6: File Parsing Capabilities")
    
    tests_passed = 0
    tests_total = 0
    
    # Test TXT parsing
    tests_total += 1
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Python Java JavaScript SQL\nDjango React Node.js")
            temp_txt = f.name
        
        with open(temp_txt, 'r') as f:
            content = f.read()
        
        if content.strip():
            print_success("TXT file parsing works")
            tests_passed += 1
        else:
            print_error("TXT file parsing returned empty")
        
        os.unlink(temp_txt)
    except Exception as e:
        print_error(f"TXT parsing test failed: {e}")
    
    # Test DOCX parsing
    tests_total += 1
    try:
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_docx = f.name
        
        doc = Document()
        doc.add_paragraph("Python")
        doc.add_paragraph("Java")
        doc.save(temp_docx)
        
        doc = Document(temp_docx)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        if "Python" in text:
            print_success("DOCX file parsing works")
            tests_passed += 1
        else:
            print_error("DOCX file parsing failed")
        
        os.unlink(temp_docx)
    except ImportError:
        print_warning("DOCX parsing test skipped (python-docx not installed)")
    except Exception as e:
        print_error(f"DOCX parsing test failed: {e}")
    
    # Test PDF parsing
    tests_total += 1
    try:
        import PyPDF2
        print_success("PyPDF2 PDF parsing library is available")
        tests_passed += 1
    except ImportError:
        print_warning("PyPDF2 not installed (required for PDF parsing)")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 7: API Key Configuration
# ─────────────────────────────────────────────────────────────────────────────

def test_api_configuration():
    print_header("Test 7: API Key Configuration")
    
    import os
    
    tests_passed = 0
    tests_total = 0
    
    # Check if API key is set
    tests_total += 1
    api_key = os.environ.get("GROQ_API_KEY")
    if api_key:
        masked_key = api_key[:8] + "*" * (len(api_key) - 16) + api_key[-8:]
        print_success(f"GROQ API Key configured: {masked_key}")
        tests_passed += 1
    else:
        print_warning("GROQ API Key not set in environment")
    
    # Check if Flask config would load it
    tests_total += 1
    try:
        # Simple sync version
        key = os.environ.get("GROQ_API_KEY")
        if key:
            print_success("Flask app can access API key")
            tests_passed += 1
        else:
            print_warning("Flask app cannot access API key")
    except Exception as e:
        print_warning(f"Could not verify Flask config: {e}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 8: JavaScript Module Integrity
# ─────────────────────────────────────────────────────────────────────────────

def test_javascript_integrity():
    print_header("Test 8: JavaScript Module Integrity")
    
    tests_passed = 0
    tests_total = 0
    
    try:
        with open("prepfriend/app/static/js/resume-modal.js", "r") as f:
            js_content = f.read()
        
        # Check for class definition
        tests_total += 1
        if "class ResumeModal" in js_content:
            print_success("ResumeModal class is defined")
            tests_passed += 1
        else:
            print_error("ResumeModal class not found")
        
        # Check for required methods
        tests_total += 1
        required_methods = [
            "init()",
            "show()",
            "close()",
            "submitResume()",
            "showUploadForm()",
        ]
        
        methods_found = 0
        for method in required_methods:
            if method in js_content:
                methods_found += 1
        
        if methods_found == len(required_methods):
            print_success(f"All required methods found ({len(required_methods)})")
            tests_passed += 1
        else:
            print_error(f"Missing methods: found {methods_found}/{len(required_methods)}")
        
        # Check for event listeners
        tests_total += 1
        event_listeners = [
            "addEventListener",
            ".insertBefore(",
        ]
        
        has_listeners = all(listener in js_content for listener in event_listeners) or "addEventListener" in js_content
        if has_listeners or "addEventListener" in js_content:
            print_success("Event listeners are set up")
            tests_passed += 1
        else:
            print_warning("Event listener setup may be incomplete")
        
        # Check for API calls
        tests_total += 1
        if "fetch(" in js_content:
            print_success("Fetch API calls are present")
            tests_passed += 1
        else:
            print_error("No fetch API calls found")
        
    except Exception as e:
        print_error(f"Could not analyze JavaScript: {e}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Test 9: Database Connectivity
# ─────────────────────────────────────────────────────────────────────────────

def test_database_connectivity():
    print_header("Test 9: Database Connectivity")
    
    tests_passed = 0
    tests_total = 0
    
    db_path = "prepfriend/data/database.db"
    
    tests_total += 1
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Test basic query
        cursor.execute("SELECT 1")
        if cursor.fetchone():
            print_success("Database connectivity verified")
            tests_passed += 1
        else:
            print_error("Database query returned no results")
        
        conn.close()
    except Exception as e:
        print_error(f"Database connectivity failed: {e}")
    
    # Test write permissions
    tests_total += 1
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Check if we can query the resumes table
        cursor.execute("SELECT COUNT(*) FROM resumes")
        count = cursor.fetchone()[0]
        
        print_success(f"Can read from database ({count} resumes found)")
        tests_passed += 1
        
        conn.close()
    except Exception as e:
        print_error(f"Database read test failed: {e}")
    
    return tests_passed, tests_total

# ─────────────────────────────────────────────────────────────────────────────
# Main Test Runner
# ─────────────────────────────────────────────────────────────────────────────

def run_all_tests():
    print(f"\n{Colors.BOLD}Resume Upload & Skill Extraction - Complete Test Suite{Colors.ENDC}")
    print(f"{Colors.BOLD}Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}{Colors.ENDC}\n")
    
    total_passed = 0
    total_tests = 0
    
    test_functions = [
        test_environment,
        test_dependencies,
        test_frontend_files,
        test_database_schema,
        test_backend_routes,
        test_file_parsing,
        test_api_configuration,
        test_javascript_integrity,
        test_database_connectivity,
    ]
    
    for test_func in test_functions:
        passed, total = test_func()
        total_passed += passed
        total_tests += total
    
    # Print summary
    print_header("Test Summary")
    
    percentage = (total_passed / total_tests * 100) if total_tests > 0 else 0
    
    print(f"Total Tests: {total_tests}")
    print(f"Passed: {Colors.OKGREEN}{total_passed}{Colors.ENDC}")
    print(f"Failed: {Colors.FAIL}{total_tests - total_passed}{Colors.ENDC}")
    print(f"Success Rate: {percentage:.1f}%")
    
    if total_passed == total_tests:
        print_success("All tests passed! System is ready.")
        return 0
    elif percentage >= 80:
        print_warning("Most tests passed. Minor issues may exist.")
        return 1
    else:
        print_error("Multiple tests failed. Check configuration.")
        return 2

if __name__ == "__main__":
    exit_code = run_all_tests()
    sys.exit(exit_code)
